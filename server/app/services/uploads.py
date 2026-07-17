import json
import subprocess
import warnings
import zipfile
from pathlib import Path

from docx import Document
from fastapi import HTTPException
from PIL import Image
from pypdf import PdfReader
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.config import Settings
from app.db.models import Asset
from app.services.assets import register_file
from app.services.workspace import project_or_404

MIB = 1024 * 1024
PROJECT_LIMIT = 1024 * MIB
TYPE_RULES: dict[str, tuple[int, str, str]] = {
    ".txt": (10 * MIB, "text/plain", "REFERENCE_TEXT"),
    ".md": (10 * MIB, "text/markdown", "REFERENCE_TEXT"),
    ".pdf": (50 * MIB, "application/pdf", "REFERENCE_DOCUMENT"),
    ".docx": (
        50 * MIB,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "REFERENCE_DOCUMENT",
    ),
    ".png": (50 * MIB, "image/png", "REFERENCE_IMAGE"),
    ".jpg": (50 * MIB, "image/jpeg", "REFERENCE_IMAGE"),
    ".jpeg": (50 * MIB, "image/jpeg", "REFERENCE_IMAGE"),
    ".webp": (50 * MIB, "image/webp", "REFERENCE_IMAGE"),
    ".mp4": (250 * MIB, "video/mp4", "REFERENCE_VIDEO"),
}


def upload_rule(filename: str) -> tuple[str, int, str, str]:
    if not filename or len(filename) > 255 or Path(filename).name != filename or "\x00" in filename:
        raise HTTPException(status_code=422, detail="素材文件名无效")
    extension = Path(filename).suffix.lower()
    rule = TYPE_RULES.get(extension)
    if rule is None:
        raise HTTPException(
            status_code=415,
            detail={
                "code": "UPLOAD_TYPE_UNSUPPORTED",
                "message": "仅支持 TXT、MD、PDF、DOCX、PNG、JPG、WebP 与 MP4",
                "retryable": False,
            },
        )
    limit, mime, kind = rule
    return extension, limit, mime, kind


def ensure_project_capacity(session: Session, project_id: str, incoming_bytes: int) -> None:
    project_or_404(session, project_id)
    current = session.scalar(
        select(func.coalesce(func.sum(Asset.size_bytes), 0)).where(Asset.project_id == project_id)
    )
    if int(current or 0) + incoming_bytes > PROJECT_LIMIT:
        raise HTTPException(
            status_code=413,
            detail={
                "code": "PROJECT_ASSET_LIMIT",
                "message": "项目素材总量不能超过 1 GB",
                "retryable": False,
            },
        )


def _read_text(path: Path) -> dict[str, object]:
    raw = path.read_bytes()
    text: str | None = None
    encoding = ""
    for candidate in ("utf-8-sig", "utf-16"):
        try:
            text = raw.decode(candidate)
            encoding = candidate
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        raise HTTPException(status_code=422, detail="TXT/MD 必须使用 UTF-8 或 UTF-16 编码")
    return {
        "parse_status": "READY",
        "encoding": encoding,
        "character_count": len(text),
        "parsed_text": text[:50_000],
        "truncated": len(text) > 50_000,
    }


def _read_pdf(path: Path) -> dict[str, object]:
    with path.open("rb") as source:
        header = source.read(5)
    if header != b"%PDF-":
        raise HTTPException(status_code=422, detail="PDF 文件头无效")
    try:
        reader = PdfReader(path, strict=True)
        if reader.is_encrypted:
            raise HTTPException(status_code=422, detail="暂不支持加密 PDF")
        pages = reader.pages[:20]
        text = "\n".join((page.extract_text() or "") for page in pages).strip()
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=422, detail="PDF 结构无效或无法解析") from exc
    return {
        "parse_status": "READY" if text else "UNSUPPORTED_OCR",
        "page_count": len(reader.pages),
        "parsed_pages": len(pages),
        "parsed_text": text[:50_000],
        "truncated": len(text) > 50_000,
    }


def _validate_docx_archive(path: Path) -> None:
    try:
        with zipfile.ZipFile(path) as archive:
            members = archive.infolist()
            if len(members) > 2_000 or sum(item.file_size for item in members) > 100 * MIB:
                raise HTTPException(status_code=422, detail="DOCX 解压规模超过安全限制")
            for item in members:
                member = Path(item.filename)
                if member.is_absolute() or ".." in member.parts:
                    raise HTTPException(status_code=422, detail="DOCX 包含不安全路径")
            if "word/document.xml" not in archive.namelist():
                raise HTTPException(status_code=422, detail="DOCX 缺少正文结构")
    except HTTPException:
        raise
    except (OSError, zipfile.BadZipFile) as exc:
        raise HTTPException(status_code=422, detail="DOCX 文件结构无效") from exc


def _read_docx(path: Path) -> dict[str, object]:
    _validate_docx_archive(path)
    try:
        document = Document(path)
        blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                blocks.append("\t".join(cell.text for cell in row.cells))
        text = "\n".join(blocks)
    except Exception as exc:
        raise HTTPException(status_code=422, detail="DOCX 无法解析") from exc
    return {
        "parse_status": "READY" if text.strip() else "EMPTY",
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "parsed_text": text[:50_000],
        "truncated": len(text) > 50_000,
    }


def _read_image(path: Path, expected_mime: str) -> tuple[dict[str, object], int, int]:
    Image.MAX_IMAGE_PIXELS = 50_000_000
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("error", Image.DecompressionBombWarning)
            with Image.open(path) as image:
                image.verify()
            with Image.open(path) as image:
                width, height = image.size
                detected_mime = Image.MIME.get(image.format or "")
    except Exception as exc:
        raise HTTPException(status_code=422, detail="图片文件头、尺寸或编码无效") from exc
    if detected_mime != expected_mime:
        raise HTTPException(status_code=422, detail="图片扩展名与文件内容不一致")
    return (
        {"parse_status": "READY", "width": width, "height": height},
        width,
        height,
    )


def _read_video(path: Path) -> tuple[dict[str, object], int | None, int | None, int]:
    with path.open("rb") as source:
        header = source.read(32)
    if len(header) < 12 or header[4:8] != b"ftyp":
        raise HTTPException(status_code=422, detail="MP4 文件头无效")
    result = subprocess.run(
        [
            "ffprobe",
            "-v",
            "error",
            "-show_entries",
            "format=duration:stream=codec_type,codec_name,width,height",
            "-of",
            "json",
            str(path),
        ],
        capture_output=True,
        text=True,
        check=False,
    )
    if result.returncode != 0:
        raise HTTPException(status_code=422, detail="MP4 无法通过 ffprobe 校验")
    probe = json.loads(result.stdout)
    streams = probe.get("streams", [])
    video = next((item for item in streams if item.get("codec_type") == "video"), None)
    if video is None:
        raise HTTPException(status_code=422, detail="MP4 缺少视频轨")
    duration_ms = round(float(probe.get("format", {}).get("duration", 0)) * 1000)
    if duration_ms <= 0:
        raise HTTPException(status_code=422, detail="MP4 时长无效")
    width = int(video.get("width") or 0) or None
    height = int(video.get("height") or 0) or None
    return (
        {
            "parse_status": "READY",
            "duration_ms": duration_ms,
            "width": width,
            "height": height,
            "streams": [
                {"type": item.get("codec_type"), "codec": item.get("codec_name")}
                for item in streams
            ],
        },
        width,
        height,
        duration_ms,
    )


def validate_and_register_upload(
    session: Session,
    settings: Settings,
    *,
    project_id: str,
    source: Path,
    filename: str,
    declared_content_type: str | None,
) -> Asset:
    extension, _limit, mime, kind = upload_rule(filename)
    size = source.stat().st_size
    ensure_project_capacity(session, project_id, size)
    width: int | None = None
    height: int | None = None
    duration_ms: int | None = None
    if extension in {".txt", ".md"}:
        metadata = _read_text(source)
    elif extension == ".pdf":
        metadata = _read_pdf(source)
    elif extension == ".docx":
        metadata = _read_docx(source)
    elif extension in {".png", ".jpg", ".jpeg", ".webp"}:
        metadata, width, height = _read_image(source, mime)
    else:
        metadata, width, height, duration_ms = _read_video(source)
    metadata.update(
        {
            "extension": extension,
            "declared_content_type": declared_content_type,
            "security_checks": ["size", "magic", "safe_path"],
        }
    )
    asset = register_file(
        session,
        settings,
        project_id=project_id,
        kind=kind,
        source=source,
        source_entity_type="project",
        source_entity_id=project_id,
        mime=mime,
        width=width,
        height=height,
        duration_ms=duration_ms,
    )
    asset.original_filename = filename
    asset.metadata_json = json.dumps(metadata, ensure_ascii=False, sort_keys=True)
    asset.rights_status = "USER_CONFIRMED"
    asset.is_temporary = False
    session.commit()
    session.refresh(asset)
    return asset
